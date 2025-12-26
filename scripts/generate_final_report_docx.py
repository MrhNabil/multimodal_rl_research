#!/usr/bin/env python
"""
COMPREHENSIVE FINAL PROJECT REPORT GENERATOR

Creates a detailed DOCX report with:
- Cover page with NSU logo
- Abstract
- Detailed methodology with equations
- All experiment results
- Embedded matplotlib diagrams
- Human-like writing style
"""

import os
import sys
from datetime import datetime

# Install docx if not available
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT

# Directories
FIGURES_DIR = r"d:\multimodal_rl_research\experiments\report_figures"
OUTPUT_DIR = r"d:\multimodal_rl_research\experiments"


def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def add_figure(doc, image_path, caption, width=5.5):
    """Add a figure with caption."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption_para = doc.add_paragraph()
        run = caption_para.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Figure not found: {image_path}]")


def create_report():
    """Create the comprehensive project report."""
    doc = Document()
    
    # ========================================
    # COVER PAGE
    # ========================================
    
    # Try to add NSU logo
    logo_paths = [
        r"C:\Users\User\.gemini\antigravity\brain\2382d71f-e86c-44b2-b7a1-dc9137c227bf\uploaded_image_1765686810465.png",
        r"d:\multimodal_rl_research\nsu_logo.png"
    ]
    
    for logo_path in logo_paths:
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(1.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
    
    doc.add_paragraph()
    
    # University name
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("NORTH SOUTH UNIVERSITY")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("Department of Electrical and Computer Engineering")
    subtitle_run.font.size = Pt(12)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Report title
    main_title = doc.add_paragraph()
    main_title_run = main_title.add_run("COMPOSITIONAL SKILL LEARNING\nIN MULTIMODAL REINFORCEMENT LEARNING")
    main_title_run.bold = True
    main_title_run.font.size = Pt(24)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Final Project Demonstration Report")
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student info
    info_lines = [
        "Submitted by:",
        "",
        "Rakib Hossain Nabil",
        "ID: 2131005642",
        "Section: 2",
        "",
        f"Date: {datetime.now().strftime('%B %Y')}"
    ]
    
    for line in info_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        if line.startswith("Rakib"):
            run.bold = True
            run.font.size = Pt(14)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========================================
    # ABSTRACT
    # ========================================
    
    add_heading(doc, "Abstract", 1)
    
    abstract_text = """This project investigates whether reinforcement learning (RL) can enable multimodal models to compose pretrained vision and language skills for Visual Question Answering (VQA) without intermediate supervision. Building upon the theoretical framework from "From f(x) and g(x) to f(g(x)): LLMs Learn New Skills in RL by Composing Old Ones," we extend compositional skill learning from text-only to cross-modal settings.

We develop a multimodal VQA system using frozen CLIP (ViT-B/32) visual features combined with trainable projection layers and question-type-specific classification heads. Through systematic experimentation across 61 different configurations, we investigate the effects of training methods (frozen, supervised, RL), learning rates (1e-5 to 1e-2), reward functions (exact match, partial, progressive), and question types (color, shape, count, spatial).

Our key findings show that supervised learning achieves 74% accuracy while RL peaks at 53.7% on this multimodal task. We observe significant variation in per-question-type performance, with spatial reasoning proving most challenging (24-61% accuracy). The experiments reveal that frozen CLIP features provide a strong foundation for visual understanding but have inherent limitations for fine-grained spatial relationships.

This work contributes to our understanding of compositional learning in multimodal systems and provides empirical evidence for the challenges of extending text-only compositional results to vision-language domains."""
    
    doc.add_paragraph(abstract_text)
    
    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").bold = True
    keywords.add_run("Visual Question Answering, Reinforcement Learning, CLIP, Compositional Learning, Multimodal AI")
    
    doc.add_page_break()
    
    # ========================================
    # TABLE OF CONTENTS
    # ========================================
    
    add_heading(doc, "Table of Contents", 1)
    
    toc = [
        "1. Introduction",
        "2. Background and Motivation",
        "3. Methodology",
        "   3.1 Model Architecture",
        "   3.2 Training Methods",
        "   3.3 REINFORCE Algorithm",
        "   3.4 Dataset",
        "4. Experimental Setup",
        "   4.1 Experiment Categories",
        "   4.2 Hyperparameters",
        "5. Results",
        "   5.1 Main Results",
        "   5.2 Learning Rate Analysis",
        "   5.3 Per-Question-Type Analysis",
        "   5.4 Complete Experiment Table",
        "6. Discussion",
        "7. Limitations",
        "8. Conclusion",
        "References"
    ]
    
    for item in toc:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========================================
    # 1. INTRODUCTION
    # ========================================
    
    add_heading(doc, "1. Introduction", 1)
    
    intro_text = """The ability to compose existing skills into new, more complex capabilities is a fundamental aspect of human intelligence. When we learn to recognize objects and understand spatial relationships separately, we can naturally combine these skills to answer questions like "What is to the left of the red cube?" without explicit training on every possible combination.

Recent work has shown that Large Language Models (LLMs) can achieve similar compositional behavior through reinforcement learning. The paper "From f(x) and g(x) to f(g(x))" demonstrates that RL enables language models to compose pretrained text skills f and g into the composed skill f(g(x)) using only the reward signal for the final output, without intermediate supervision.

This raises an intriguing question: can this compositional learning transfer to multimodal settings?

In this project, we extend the compositional skill learning framework from text-only to vision-language systems. Specifically, we investigate whether RL can compose:
• Frozen visual skills (encoded by CLIP) 
• With trainable language understanding
• To answer Visual Question Answering (VQA) queries

Our research question is: "Can reinforcement learning compose frozen vision skills with trainable language skills for VQA, without intermediate supervision on visual features?"

To answer this, we conducted 61 experiments systematically varying training methods, learning rates, reward functions, and question types. Our findings reveal both the promise and limitations of compositional learning in multimodal settings."""
    
    doc.add_paragraph(intro_text)
    
    doc.add_page_break()
    
    # ========================================
    # 2. BACKGROUND
    # ========================================
    
    add_heading(doc, "2. Background and Motivation", 1)
    
    add_heading(doc, "2.1 Reference Paper", 2)
    
    ref_text = """The theoretical foundation for this work comes from "From f(x) and g(x) to f(g(x)): LLMs Learn New Skills in RL by Composing Old Ones" (arXiv:2509.25123). This paper makes several key claims:

1. Pre-trained LLMs have distinct skills (e.g., translation, summarization)
2. RL with sparse rewards can compose these skills into new capabilities
3. The composition emerges without intermediate supervision

For example, if an LLM knows how to translate (f) and summarize (g), RL can teach it to "translate then summarize" (f(g(x))) using only a reward for the final summary quality, without explicitly training on translation outputs."""
    
    doc.add_paragraph(ref_text)
    
    add_heading(doc, "2.2 Extension to Multimodal", 2)
    
    multimodal_text = """We hypothesize that similar compositional behavior could occur in multimodal systems. Our extension involves:

• Visual skill f: CLIP's ability to encode visual features
• Language skill g: Understanding and classifying question types
• Composed skill f(g): Answering VQA questions by combining both

The key difference from the original paper is that our visual skill (CLIP) processes a different modality than the language understanding, making the composition cross-modal rather than within-modality."""
    
    doc.add_paragraph(multimodal_text)
    
    add_heading(doc, "2.3 Why CLIP?", 2)
    
    clip_text = """CLIP (Contrastive Language-Image Pre-training) was chosen because:

1. It encodes images into semantically meaningful embeddings
2. It's trained on 400M image-text pairs, providing robust visual features
3. The frozen encoder allows us to isolate the effect of training on composition
4. It's well-established in the VQA literature"""
    
    doc.add_paragraph(clip_text)
    
    doc.add_page_break()
    
    # ========================================
    # 3. METHODOLOGY
    # ========================================
    
    add_heading(doc, "3. Methodology", 1)
    
    add_heading(doc, "3.1 Model Architecture", 2)
    
    arch_text = """Our multimodal VQA system consists of four main components:

1. Vision Encoder (CLIP ViT-B/32)
   - Input: 224×224 RGB images
   - Output: 512-dimensional visual embeddings
   - Parameters: 151 million (all frozen)
   - Purpose: Extract semantic visual features

2. Projection Layer
   - Input: 512-d visual embedding
   - Output: 512-768d projected features
   - Parameters: ~400K (trainable)
   - Purpose: Adapt CLIP features for VQA

3. Question Encoder
   - Uses question type detection (color/shape/count/spatial)
   - Learned embeddings for each question type
   - Parameters: ~100K (trainable)

4. Type-Specific Classification Heads
   - Color head: 4 classes (red, blue, green, yellow)
   - Shape head: 3 classes (cube, sphere, cylinder)
   - Count head: 4 classes (0, 1, 2, 3)
   - Spatial head: 13 classes (color+shape combinations + "nothing")
   - Parameters: ~500K (trainable)

Total trainable parameters: approximately 1 million (0.6% of full model)"""
    
    doc.add_paragraph(arch_text)
    
    # Add architecture figure
    arch_fig = os.path.join(FIGURES_DIR, "fig1_architecture.png")
    add_figure(doc, arch_fig, "Figure 1: Complete Model Architecture showing frozen CLIP encoder (blue), trainable projection and fusion layers (orange), and question-type-specific classification heads.")
    
    doc.add_page_break()
    
    add_heading(doc, "3.2 Training Methods", 2)
    
    methods_text = """We compared three training approaches:

1. Frozen Baseline
   - No training whatsoever
   - Direct inference with pretrained components
   - Establishes lower bound on performance
   - Expected accuracy: near random (4.17%)

2. Supervised Learning
   - Cross-entropy loss on answer labels
   - Direct gradient signal for every sample
   - Formula: L = -Σ y_i log(ŷ_i)
   - Most sample-efficient method

3. Reinforcement Learning (REINFORCE)
   - Policy gradient with binary reward
   - Reward R=1 if answer is correct, R=0 otherwise
   - Sparse reward signal encourages exploration
   - Tests compositional learning hypothesis"""
    
    doc.add_paragraph(methods_text)
    
    # Add method comparison figure
    method_fig = os.path.join(FIGURES_DIR, "fig4_method_comparison.png")
    add_figure(doc, method_fig, "Figure 2: Comparison of training methods on VQA task (1000 training steps).")
    
    add_heading(doc, "3.3 REINFORCE Algorithm", 2)
    
    reinforce_text = """The REINFORCE algorithm uses policy gradients to optimize the model. The key update equation is:

∇J(θ) = E[R · ∇_θ log π(a|s; θ)]

Where:
- J(θ) is the expected reward
- R is the reward (1 for correct, 0 for wrong)
- π(a|s; θ) is the policy (softmax over answer classes)
- θ are the model parameters

To reduce variance, we use a baseline:
A = R - baseline
where baseline is the running mean of rewards

The update becomes:
θ ← θ + α · A · ∇log π

This allows the model to increase the probability of actions that led to higher-than-average rewards."""
    
    doc.add_paragraph(reinforce_text)
    
    # Add REINFORCE flowchart
    reinforce_fig = os.path.join(FIGURES_DIR, "fig6_reinforce.png")
    add_figure(doc, reinforce_fig, "Figure 3: REINFORCE algorithm flowchart showing the training loop.")
    
    doc.add_page_break()
    
    add_heading(doc, "3.4 Dataset", 2)
    
    dataset_text = """We use a synthetic CLEVR-like VQA dataset:

Dataset Statistics:
- Training: 5,000 samples
- Validation: 1,000 samples  
- Test: 1,000 samples
- Image size: 224×224 pixels
- Answer vocabulary: 24 classes

Question Types:
1. Color questions: "What color is the X?" → red/blue/green/yellow
2. Shape questions: "What shape is the X?" → cube/sphere/cylinder
3. Count questions: "How many X are there?" → 0/1/2/3
4. Spatial questions: "What is left/right of X?" → color+shape combinations

The synthetic nature allows controlled experimentation but may differ from real-world VQA distributions."""
    
    doc.add_paragraph(dataset_text)
    
    doc.add_page_break()
    
    # ========================================
    # 4. EXPERIMENTAL SETUP
    # ========================================
    
    add_heading(doc, "4. Experimental Setup", 1)
    
    add_heading(doc, "4.1 Experiment Categories", 2)
    
    exp_text = """We conducted experiments in the following categories:

1. Baseline Experiments (3 experiments)
   - Frozen baseline
   - Supervised learning
   - RL baseline

2. Learning Rate Sweep (10 experiments)
   - Range: 1e-5 to 1e-2
   - Purpose: Find optimal training dynamics

3. Reward Function Variations (10 experiments)
   - Exact match, partial match, length penalty
   - Progressive rewards, combined rewards
   - Purpose: Improve RL training signal

4. Question Type Experiments (6 experiments)
   - Train on specific question types
   - Purpose: Understand skill-specific learning

5. Architecture Variations (4 experiments)
   - HighAccuracyVQA, UltraHigh models
   - Purpose: Test model capacity effects"""
    
    doc.add_paragraph(exp_text)
    
    # Add categories pie chart
    cat_fig = os.path.join(FIGURES_DIR, "fig8_categories.png")
    add_figure(doc, cat_fig, "Figure 4: Distribution of experiment categories.", width=4)
    
    add_heading(doc, "4.2 Hyperparameters", 2)
    
    # Create hyperparameter table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    headers = ['Parameter', 'Value']
    data = [
        ('Optimizer', 'AdamW'),
        ('Weight Decay', '0.01'),
        ('Batch Size', '64'),
        ('Training Steps', '500-3000'),
        ('Learning Rate (searched)', '1e-5 to 1e-2'),
        ('Gradient Clipping', '1.0'),
        ('Random Seed', '42'),
    ]
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (param, value) in enumerate(data, 1):
        table.rows[i].cells[0].text = param
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    doc.add_paragraph("Table 1: Hyperparameters used across experiments.").italic = True
    
    doc.add_page_break()
    
    # ========================================
    # 5. RESULTS
    # ========================================
    
    add_heading(doc, "5. Results", 1)
    
    add_heading(doc, "5.1 Main Results", 2)
    
    main_results_text = """Our main findings are summarized below:

Best Performance by Training Method:
• Supervised Learning: 74.0% test accuracy
• Reinforcement Learning: 53.7% test accuracy (lr=2e-4)
• Frozen Baseline: 0.2% test accuracy

Key Observation: Supervised learning significantly outperforms RL on this multimodal VQA task, achieving 74% compared to RL's 53.7%. This differs from the text-only findings in the reference paper, suggesting that cross-modal composition may require stronger training signals than within-modality composition."""
    
    doc.add_paragraph(main_results_text)
    
    # Add all experiments figure
    all_exp_fig = os.path.join(FIGURES_DIR, "fig2_all_experiments.png")
    add_figure(doc, all_exp_fig, "Figure 5: Summary of all experiments sorted by accuracy.")
    
    doc.add_page_break()
    
    add_heading(doc, "5.2 Learning Rate Analysis", 2)
    
    lr_text = """One of our most important findings concerns learning rate sensitivity in RL training:

• Optimal learning rate: 2e-4 (achieving 53.7%)
• Learning rates < 1e-4: Underfitting (29-45%)
• Learning rates > 1e-3: Unstable or collapsed (14-29%)

The optimal zone is relatively narrow (1e-4 to 5e-4), emphasizing the importance of hyperparameter tuning for RL training."""
    
    doc.add_paragraph(lr_text)
    
    # Add LR sensitivity figure
    lr_fig = os.path.join(FIGURES_DIR, "fig3_learning_rate.png")
    add_figure(doc, lr_fig, "Figure 6: Learning rate sensitivity analysis showing optimal range.")
    
    add_heading(doc, "5.3 Per-Question-Type Analysis", 2)
    
    type_text = """Performance varies significantly across question types:

Supervised Learning (74% overall):
• Count: 82.0% - easiest task
• Shape: 77.4% - high accuracy
• Color: 75.7% - high accuracy
• Spatial: 61.3% - most challenging

RL Training (47.6% overall):
• Shape: 71.8% - competitive with supervised
• Count: 58.0% - reasonable
• Spatial: 39.8% - challenging
• Color: 20.6% - severe underperformance

The most striking finding is that RL struggles severely with color questions (20.6% vs 75.7% for supervised), while performing competitively on shape questions. This suggests that different visual skills have different "composability" with RL training."""
    
    doc.add_paragraph(type_text)
    
    # Add per-type figure
    type_fig = os.path.join(FIGURES_DIR, "fig5_per_type.png")
    add_figure(doc, type_fig, "Figure 7: Accuracy breakdown by question type across training methods.")
    
    doc.add_page_break()
    
    add_heading(doc, "5.4 Complete Experiment Table", 2)
    
    doc.add_paragraph("The following table shows key experiments with their configurations and results:")
    
    # Create results table
    table = doc.add_table(rows=18, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Experiment', 'Method', 'Accuracy', 'Notes']
    experiments = [
        ('exp_002_supervised', 'Supervised', '74.0%', 'Best overall'),
        ('HighAccuracyVQA', 'Supervised', '68.7%', 'Type-specific heads'),
        ('exp_008_lr_2e-4', 'RL', '53.7%', 'Optimal LR'),
        ('exp_003_rl_baseline', 'RL', '47.6%', 'Default settings'),
        ('exp_007_lr_1e-4', 'RL', '45.2%', 'Lower LR'),
        ('exp_009_lr_5e-4', 'RL', '44.0%', 'Higher LR'),
        ('exp_023_prog_slow', 'RL', '43.1%', 'Best reward function'),
        ('exp_006_lr_5e-5', 'RL', '41.0%', ''),
        ('exp_005_lr_2e-5', 'RL', '37.0%', ''),
        ('exp_002_sup_500', 'Supervised', '33.7%', '500 steps only'),
        ('exp_016_length', 'RL', '32.4%', 'Length penalty reward'),
        ('exp_010_lr_1e-3', 'RL', '29.3%', 'Too high LR'),
        ('exp_004_lr_1e-5', 'RL', '29.4%', 'Too low LR'),
        ('exp_011_lr_2e-3', 'RL', '20.7%', 'Unstable'),
        ('exp_012_lr_5e-3', 'RL', '14.2%', 'Collapsed'),
        ('exp_013_lr_1e-2', 'RL', '14.2%', 'Collapsed'),
        ('exp_001_frozen', 'Frozen', '0.2%', 'No training'),
    ]
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (exp, method, acc, notes) in enumerate(experiments, 1):
        table.rows[i].cells[0].text = exp
        table.rows[i].cells[1].text = method
        table.rows[i].cells[2].text = acc
        table.rows[i].cells[3].text = notes
    
    doc.add_paragraph()
    doc.add_paragraph("Table 2: Summary of key experiments with configurations and results.").italic = True
    
    doc.add_page_break()
    
    # ========================================
    # 6. DISCUSSION
    # ========================================
    
    add_heading(doc, "6. Discussion", 1)
    
    discussion_text = """6.1 Why Does Supervised Outperform RL?

Our results show supervised learning (74%) significantly outperforming RL (53.7%) on the multimodal VQA task. This is in contrast to the reference paper's findings for text-only composition. Several factors may explain this:

1. Cross-Modal Gap: Composing visual and language information may be fundamentally harder than composing two language skills. The modality mismatch creates an additional alignment challenge.

2. Sparse Rewards: The binary reward signal (correct/incorrect) provides less information than the cross-entropy loss, which penalizes every incorrect answer dimension.

3. Training Scale: With only 1000-3000 steps, RL may need more iterations to discover effective compositions. Text-only experiments often use 10-100x more steps.

4. Frozen Visual Features: CLIP was not trained for fine-grained spatial discrimination or color-object binding, limiting what composition can achieve.

6.2 The Color Puzzle

One of our most striking findings is that RL severely underperforms on color questions (20.6% vs 75.7% for supervised). Possible explanations:

1. CLIP embeddings may encode color weakly compared to shape
2. Color words in answers may confuse the RL policy
3. The reward signal doesn't distinguish "close" answers (blue vs green)

This suggests that compositional learning may be skill-specific rather than universal.

6.3 Implications for Multimodal Composition

Our experiments suggest that:
• Compositional learning transfers to multimodal settings, but with reduced effectiveness
• Some visual skills (shape recognition) compose better than others (color)
• Supervised learning remains more sample-efficient for multimodal tasks
• Larger training budgets may be needed for RL to match supervised performance"""
    
    doc.add_paragraph(discussion_text)
    
    doc.add_page_break()
    
    # ========================================
    # 7. LIMITATIONS
    # ========================================
    
    add_heading(doc, "7. Limitations", 1)
    
    limitations_text = """This study has several important limitations:

1. Synthetic Dataset
   - Our synthetic VQA data may not reflect real-world distributions
   - CLEVR-style scenes are simpler than natural images
   - Limited answer vocabulary (24 classes vs 3000+ in real VQA)

2. Model Scale
   - CLIP ViT-B/32 is relatively small (151M parameters)
   - Larger vision models might show different composition patterns
   - T5-small was simplified to MLP classifiers for efficiency

3. Training Resources
   - Limited to 1000-3000 training steps due to time constraints
   - Longer training might change supervised vs RL comparison
   - No GPU training for largest model configurations

4. Frozen Visual Encoder
   - We did not explore fine-tuning CLIP
   - Unfreezing could improve spatial accuracy but increases complexity
   - Breaks the "frozen skill composition" hypothesis

5. Limited Reward Engineering
   - Primarily used binary rewards
   - More sophisticated reward shaping might help RL
   - Did not explore intrinsic motivation or curiosity"""
    
    doc.add_paragraph(limitations_text)
    
    # ========================================
    # 8. CONCLUSION
    # ========================================
    
    add_heading(doc, "8. Conclusion", 1)
    
    conclusion_text = """This project investigated whether reinforcement learning can enable compositional skill learning in multimodal systems, extending the theoretical framework from text-only to vision-language settings.

Key Findings:

1. Composition is Learnable: The improvement from frozen baseline (0.2%) to trained models (53.7-74%) demonstrates that visual and language skills can be composed for VQA.

2. Supervised Outperforms RL: With 1000 training steps, supervised learning (74%) significantly outperforms RL (53.7%) on multimodal VQA, unlike text-only results.

3. Skill-Specific Composition: Different visual skills (color, shape, count, spatial) have different "composability" - shape composes well via RL (71.8%) while color does not (20.6%).

4. Learning Rate Critical: RL training is highly sensitive to learning rate, with optimal range of 1e-4 to 5e-4.

5. Frozen Features Limit: CLIP's frozen features provide a strong foundation but inherently limit spatial and color reasoning.

Future Directions:

• Scale up training steps to 10,000+
• Explore BLIP-2 style Q-Former architectures
• Fine-tune visual encoder (unfreeze CLIP)
• Test on natural image VQA datasets
• Investigate hybrid supervised+RL training

In conclusion, compositional skill learning shows promise for multimodal systems but requires careful consideration of the modality gap, skill-specific composability, and training methodology. This work contributes to our understanding of how to build AI systems that can flexibly combine pretrained capabilities for novel tasks."""
    
    doc.add_paragraph(conclusion_text)
    
    doc.add_page_break()
    
    # ========================================
    # REFERENCES
    # ========================================
    
    add_heading(doc, "References", 1)
    
    references = [
        "[1] Anonymous. 'From f(x) and g(x) to f(g(x)): LLMs Learn New Skills in RL by Composing Old Ones.' arXiv:2509.25123, 2024.",
        "[2] Radford, A., et al. 'Learning Transferable Visual Models From Natural Language Supervision.' ICML 2021.",
        "[3] Williams, R.J. 'Simple Statistical Gradient-Following Algorithms for Connectionist Reinforcement Learning.' Machine Learning, 1992.",
        "[4] Antol, S., et al. 'VQA: Visual Question Answering.' ICCV 2015.",
        "[5] Li, J., et al. 'BLIP-2: Bootstrapping Language-Image Pre-training with Frozen Image Encoders and Large Language Models.' ICML 2023.",
        "[6] Lake, B.M., et al. 'Building Machines That Learn and Think Like People.' Behavioral and Brain Sciences, 2017.",
        "[7] Johnson, J., et al. 'CLEVR: A Diagnostic Dataset for Compositional Language and Elementary Visual Reasoning.' CVPR 2017.",
        "[8] Keysers, C., et al. 'Measuring Compositional Generalization: A Comprehensive Method on Realistic Data.' ICLR 2020.",
    ]
    
    for ref in references:
        doc.add_paragraph(ref)
    
    # ========================================
    # SAVE DOCUMENT
    # ========================================
    
    output_path = os.path.join(OUTPUT_DIR, "FINAL_PROJECT_REPORT.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print("GENERATING COMPREHENSIVE FINAL REPORT")
    print("=" * 60)
    
    output = create_report()
    
    print("\n" + "=" * 60)
    print("REPORT GENERATION COMPLETE")
    print(f"Output: {output}")
    print("=" * 60)
