#!/usr/bin/env python
"""
ENHANCED FINAL REPORT GENERATOR (V2)

Google Docs compatible DOCX with:
- 10 professional figures
- More visual elements
- Cleaner formatting
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

FIGURES_DIR = r"d:\multimodal_rl_research\experiments\report_figures_v2"
OUTPUT_DIR = r"d:\multimodal_rl_research\experiments"


def add_figure(doc, image_path, caption, width=6):
    """Add figure with caption."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def create_enhanced_report():
    """Create enhanced report with more visuals."""
    doc = Document()
    
    # ==================== COVER PAGE ====================
    
    logo_path = r"C:\Users\User\.gemini\antigravity\brain\2382d71f-e86c-44b2-b7a1-dc9137c227bf\uploaded_image_1765686810465.png"
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # University
    p = doc.add_paragraph()
    r = p.add_run("NORTH SOUTH UNIVERSITY")
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    r = p.add_run("Department of Electrical and Computer Engineering")
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    r = p.add_run("COMPOSITIONAL SKILL LEARNING IN\nMULTIMODAL REINFORCEMENT LEARNING")
    r.bold = True
    r.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    r = p.add_run("Final Project Demonstration Report")
    r.italic = True
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Student info
    info = ["Submitted by:", "", "Rakib Hossain Nabil", "ID: 2131005642", "Section: 2", "", f"Date: {datetime.now().strftime('%B %Y')}"]
    for line in info:
        p = doc.add_paragraph()
        r = p.add_run(line)
        if "Rakib" in line:
            r.bold = True
            r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== ABSTRACT ====================
    
    doc.add_heading("Abstract", 1)
    
    doc.add_paragraph("""This project investigates compositional skill learning in multimodal systems, extending the theoretical framework from "From f(x) and g(x) to f(g(x))" to vision-language settings. We develop a Visual Question Answering (VQA) system using frozen CLIP visual features combined with trainable projection layers.

Through 61+ experiments, we systematically study training methods (frozen, supervised, RL), learning rates (1e-5 to 1e-2), reward functions, and question types. Our key findings show that supervised learning achieves 74% accuracy while RL peaks at 53.7%, with significant variation across question types (color: 20-76%, shape: 72-77%, count: 58-82%, spatial: 24-61%).

This work contributes empirical evidence for the challenges of extending compositional learning to multimodal domains.""")
    
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("Visual Question Answering, REINFORCE, CLIP, Multimodal AI, Compositional Learning")
    
    doc.add_page_break()
    
    # ==================== KEY FINDINGS (VISUAL) ====================
    
    doc.add_heading("Key Findings at a Glance", 1)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig10_key_findings.png"),
               "Figure 1: Summary of key research findings from 61+ experiments", 6)
    
    doc.add_page_break()
    
    # ==================== INTRODUCTION ====================
    
    doc.add_heading("1. Introduction", 1)
    
    doc.add_paragraph("""The ability to compose existing skills into new capabilities is fundamental to human intelligence. When we learn to recognize objects and understand spatial relationships separately, we can naturally combine these skills to answer questions like "What is to the left of the red cube?"

Recent work has shown that LLMs can achieve compositional behavior through reinforcement learning. The paper "From f(x) and g(x) to f(g(x))" demonstrates that RL enables models to compose pretrained skills using only sparse reward signals.

This raises our core research question:""")
    
    p = doc.add_paragraph()
    r = p.add_run("Can reinforcement learning compose frozen vision skills with trainable language skills for VQA, without intermediate supervision?")
    r.italic = True
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("""To answer this, we conducted 61+ experiments systematically varying training methods, learning rates, reward functions, and question types.""")
    
    doc.add_page_break()
    
    # ==================== METHODOLOGY ====================
    
    doc.add_heading("2. Methodology", 1)
    
    doc.add_heading("2.1 Model Architecture", 2)
    
    doc.add_paragraph("""Our multimodal VQA system uses a frozen CLIP visual encoder with trainable projection and classification layers:""")
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig01_architecture.png"),
               "Figure 2: Complete model architecture showing frozen CLIP (blue), trainable layers (orange), and output heads", 6.5)
    
    doc.add_paragraph("""Key components:
• Vision Encoder: CLIP ViT-B/32 (frozen, 151M parameters)
• Projection Layer: MLP (trainable, ~400K parameters)  
• Fusion Layer: Concatenation + MLP
• Classification Heads: 4 specialized heads for color (4), shape (3), count (4), spatial (13) classes""")
    
    doc.add_heading("2.2 Training Pipeline", 2)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig02_training_pipeline.png"),
               "Figure 3: Training pipeline showing the iterative process from data loading to evaluation", 6.5)
    
    doc.add_heading("2.3 Training Methods", 2)
    
    doc.add_paragraph("""We compared three training approaches:

1. Frozen Baseline - No training, direct inference
2. Supervised Learning - Cross-entropy loss with direct labels
3. Reinforcement Learning - REINFORCE with binary reward (R=1 if correct)""")
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig03_method_comparison.png"),
               "Figure 4: Comparison of training methods on the same dataset (1000 steps)", 5.5)
    
    doc.add_page_break()
    
    doc.add_heading("2.4 REINFORCE Algorithm", 2)
    
    doc.add_paragraph("""For RL training, we use the REINFORCE policy gradient algorithm:""")
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig07_reinforce.png"),
               "Figure 5: REINFORCE algorithm flowchart with the key policy gradient equation", 6)
    
    doc.add_paragraph("""The policy gradient update is:
∇J(θ) = E[R · ∇log π(a|s; θ)]

Where R=1 if the predicted answer matches the ground truth, R=0 otherwise. A running mean baseline reduces variance.""")
    
    doc.add_heading("2.5 Dataset", 2)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig09_dataset.png"),
               "Figure 6: Dataset statistics showing train/val/test split and balanced question types", 6)
    
    doc.add_paragraph("""We use a synthetic CLEVR-like VQA dataset:
• 7,000 total samples (5,000 train, 1,000 val, 1,000 test)
• 224×224 pixel images with colored geometric shapes
• 4 question types: Color, Shape, Count, Spatial
• 24 possible answers across all question types""")
    
    doc.add_page_break()
    
    # ==================== EXPERIMENTS ====================
    
    doc.add_heading("3. Experimental Results", 1)
    
    doc.add_heading("3.1 All Experiments Overview", 2)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig06_all_experiments.png"),
               "Figure 7: Summary of all key experiments sorted by accuracy", 6)
    
    doc.add_heading("3.2 Learning Rate Analysis", 2)
    
    doc.add_paragraph("""Learning rate is critical for RL training. We swept from 1e-5 to 1e-2:""")
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig04_learning_rate.png"),
               "Figure 8: Learning rate sensitivity showing optimal zone (1e-4 to 5e-4)", 6)
    
    doc.add_paragraph("""Key findings:
• Optimal LR: 2e-4 (53.7% accuracy)
• Too low (< 1e-4): Underfitting, slow learning
• Too high (> 1e-3): Training collapse, unstable gradients""")
    
    doc.add_heading("3.3 Per-Question-Type Performance", 2)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig05_per_type.png"),
               "Figure 9: Accuracy breakdown by question type showing RL's struggle with color", 6)
    
    doc.add_paragraph("""Critical observation: RL achieves only 20.6% on color questions compared to 75.7% for supervised learning. This suggests that some visual skills are more "composable" than others.""")
    
    doc.add_heading("3.4 Reward Function Comparison", 2)
    
    add_figure(doc, os.path.join(FIGURES_DIR, "fig08_reward_functions.png"),
               "Figure 10: Comparison of different reward functions for RL training", 5.5)
    
    doc.add_page_break()
    
    doc.add_heading("3.5 Complete Experiment Summary Table", 2)
    
    # Table of experiments
    table = doc.add_table(rows=18, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Experiment', 'Method', 'Accuracy', 'Configuration']
    experiments = [
        ('Supervised (1000 steps)', 'Supervised', '74.0%', 'lr=0.0002, 1000 steps'),
        ('HighAccuracyVQA', 'Supervised', '68.7%', 'type-specific heads'),
        ('50K Dataset', 'Supervised', '61.5%', '50K training samples'),
        ('RL lr=2e-4 (best)', 'RL', '53.7%', '3000 steps'),
        ('RL baseline', 'RL', '47.6%', 'lr=0.0002, 1000 steps'),
        ('RL lr=1e-4', 'RL', '45.2%', '1000 steps'),
        ('RL lr=5e-4', 'RL', '44.0%', '3000 steps'),
        ('RL prog_slow reward', 'RL', '43.1%', 'progressive reward'),
        ('RL lr=5e-5', 'RL', '41.0%', '1000 steps'),
        ('RL lr=2e-5', 'RL', '37.0%', '1000 steps'),
        ('Supervised (500 steps)', 'Supervised', '33.7%', '500 steps only'),
        ('RL combined reward', 'RL', '32.4%', 'multiple reward types'),
        ('RL lr=1e-3', 'RL', '29.3%', 'high LR'),
        ('RL lr=1e-5', 'RL', '29.4%', 'low LR'),
        ('RL lr=2e-3', 'RL', '20.7%', 'unstable'),
        ('RL lr=5e-3+', 'RL', '14.2%', 'collapsed'),
        ('Frozen baseline', 'Frozen', '0.2%', 'no training'),
    ]
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, (exp, method, acc, config) in enumerate(experiments, 1):
        table.rows[i].cells[0].text = exp
        table.rows[i].cells[1].text = method
        table.rows[i].cells[2].text = acc
        table.rows[i].cells[3].text = config
    
    doc.add_paragraph()
    p = doc.add_paragraph("Table 1: Summary of key experiments with configurations")
    p.runs[0].italic = True
    
    doc.add_page_break()
    
    # ==================== DISCUSSION ====================
    
    doc.add_heading("4. Discussion", 1)
    
    doc.add_heading("4.1 Why Does Supervised Outperform RL?", 2)
    
    doc.add_paragraph("""Several factors explain the 74% vs 53.7% gap:

1. Cross-Modal Composition is Harder: Composing visual and language information across modalities is more challenging than within-modality composition.

2. Sparse Rewards: Binary rewards (correct/incorrect) provide less information than cross-entropy loss, which penalizes every dimension.

3. Training Scale: With only 1000-3000 steps, RL may need 10-100x more iterations to match supervised performance.

4. Frozen Features: CLIP wasn't trained for fine-grained VQA, limiting achievable accuracy.""")
    
    doc.add_heading("4.2 The Color Puzzle", 2)
    
    doc.add_paragraph("""RL achieves only 20.6% on color questions vs 75.7% for supervised (vs 72% on shape). Possible explanations:

• CLIP may encode color weakly compared to shape
• Color words in answers may confuse the RL policy
• The reward signal doesn't distinguish "close" answers (blue vs green)

This suggests compositional learning may be skill-specific.""")
    
    doc.add_heading("4.3 Why More Data Didn't Help", 2)
    
    doc.add_paragraph("""We trained on 50K samples (10x original) but achieved only 61.5% (vs 68.7% on 5K). This counter-intuitive result suggests:

• The bottleneck is the frozen CLIP features, not data quantity
• Synthetic data may have diminishing returns
• The model capacity is saturated""")
    
    doc.add_page_break()
    
    # ==================== LIMITATIONS ====================
    
    doc.add_heading("5. Limitations", 1)
    
    doc.add_paragraph("""1. Synthetic Dataset: CLEVR-style scenes are simpler than natural images

2. Model Scale: CLIP ViT-B/32 is relatively small (151M parameters)

3. Training Steps: Limited to 1000-3000 steps due to time constraints

4. Frozen Encoder: We did not explore fine-tuning CLIP

5. Limited Reward Engineering: Primarily used binary rewards""")
    
    # ==================== CONCLUSION ====================
    
    doc.add_heading("6. Conclusion", 1)
    
    doc.add_paragraph("""This project investigated compositional skill learning in multimodal systems through 61+ experiments.

Key Findings:
• Composition is learnable: 0.2% → 74% with training
• Supervised outperforms RL: 74% vs 53.7% on multimodal VQA
• Skill-specific composition: Shape (72%) composes better than color (20%) via RL
• Learning rate critical: Optimal zone 1e-4 to 5e-4
• More data ≠ better: 50K → 61.5% < 5K → 68.7%

Future Directions:
• Scale training to 10,000+ steps
• Fine-tune CLIP encoder
• Test on natural image VQA datasets
• Explore hybrid supervised+RL training

In conclusion, compositional skill learning shows promise for multimodal systems but requires careful consideration of modality gaps and training methodology.""")
    
    doc.add_page_break()
    
    # ==================== REFERENCES ====================
    
    doc.add_heading("References", 1)
    
    refs = [
        "[1] Anonymous. 'From f(x) and g(x) to f(g(x)): LLMs Learn New Skills in RL.' arXiv:2509.25123, 2024.",
        "[2] Radford et al. 'Learning Transferable Visual Models From Natural Language Supervision.' ICML 2021.",
        "[3] Williams, R.J. 'Simple Statistical Gradient-Following Algorithms for RL.' Machine Learning, 1992.",
        "[4] Antol et al. 'VQA: Visual Question Answering.' ICCV 2015.",
        "[5] Li et al. 'BLIP-2: Bootstrapping Language-Image Pre-training.' ICML 2023.",
        "[6] Johnson et al. 'CLEVR: A Diagnostic Dataset for Compositional Reasoning.' CVPR 2017.",
        "[7] Lake et al. 'Building Machines That Learn and Think Like People.' BBS, 2017.",
        "[8] Keysers et al. 'Measuring Compositional Generalization.' ICLR 2020.",
    ]
    
    for ref in refs:
        doc.add_paragraph(ref)
    
    # ==================== SAVE ====================
    
    output_path = os.path.join(OUTPUT_DIR, "FINAL_PROJECT_REPORT_ENHANCED.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print("GENERATING ENHANCED FINAL REPORT (V2)")
    print("=" * 60)
    
    output = create_enhanced_report()
    
    print("\n" + "=" * 60)
    print("REPORT COMPLETE")
    print(f"Output: {output}")
    print("=" * 60)
