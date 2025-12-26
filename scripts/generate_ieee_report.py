#!/usr/bin/env python
"""
IEEE Two-Column Format DOCX Report Generator

Creates a professional IEEE-style report with:
- Best results first
- Two-column layout (simulated)
- Tables and figures
- References
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

FIGURES_DIR = r"d:\multimodal_rl_research\experiments\report_figures_v2"
OUTPUT_DIR = r"d:\multimodal_rl_research\experiments"


def set_narrow_margins(doc):
    """Set narrow margins for IEEE style."""
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_ieee_title(doc, title, authors):
    """Add IEEE-style title block."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Author
    for author in authors:
        p = doc.add_paragraph()
        r = p.add_run(author)
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(doc, title, level=1):
    """Add IEEE-style section heading."""
    if level == 1:
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.size = Pt(12)
        heading.runs[0].font.bold = True
    else:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.size = Pt(11)
        heading.runs[0].font.italic = True


def add_body_text(doc, text):
    """Add body text with IEEE formatting."""
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(6)
    for run in para.runs:
        run.font.size = Pt(10)
    return para


def create_ieee_report():
    """Create IEEE-style report."""
    doc = Document()
    set_narrow_margins(doc)
    
    # ========================================
    # TITLE
    # ========================================
    add_ieee_title(doc, 
        "Compositional Skill Learning in Multimodal\nReinforcement Learning for Visual Question Answering",
        [
            "Rakib Hossain Nabil",
            "Department of Electrical and Computer Engineering",
            "North South University, Dhaka, Bangladesh",
            "ID: 2131005642, Section: 2"
        ]
    )
    
    doc.add_paragraph()
    
    # ========================================
    # ABSTRACT
    # ========================================
    abstract_heading = doc.add_paragraph()
    abstract_heading.add_run("Abstract—").bold = True
    abstract_heading.add_run("""We investigate compositional skill learning in multimodal systems by extending reinforcement learning (RL) approaches from text-only to vision-language settings. Using a Visual Question Answering (VQA) task with frozen CLIP visual features and trainable classification heads, we conduct 61+ experiments across training methods, learning rates, and reward functions. Our best results achieve 74.0% accuracy with supervised learning and 53.7% with REINFORCE policy gradient. We find that RL performance varies significantly by question type: shape recognition achieves 71.8% while color questions reach only 20.6%. The optimal learning rate for RL is 2e-4, with higher rates causing training collapse. Our findings suggest that cross-modal skill composition is more challenging than within-modality composition.""")
    
    keywords = doc.add_paragraph()
    keywords.add_run("Keywords—").bold = True
    keywords.add_run("Visual Question Answering, Reinforcement Learning, CLIP, Multimodal AI, REINFORCE")
    
    doc.add_paragraph()
    
    # ========================================
    # I. BEST RESULTS SUMMARY (FIRST!)
    # ========================================
    add_section(doc, "I. BEST RESULTS SUMMARY")
    
    add_body_text(doc, "Before detailing our methodology, we present our key experimental findings. These results are from 61+ controlled experiments on a synthetic VQA dataset.")
    
    # Best Results Table
    doc.add_paragraph().add_run("Table I: Best Results by Training Method").bold = True
    
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers1 = ['Method', 'Accuracy', 'Configuration']
    data1 = [
        ('Supervised Learning', '74.0%', 'lr=2e-4, 1000 steps'),
        ('HighAccuracyVQA', '68.7%', 'type-specific heads'),
        ('RL (REINFORCE)', '53.7%', 'lr=2e-4, 3000 steps'),
        ('RL Baseline', '47.6%', 'lr=2e-4, 1000 steps'),
        ('Frozen Baseline', '0.2%', 'no training'),
    ]
    
    for i, h in enumerate(headers1):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, row in enumerate(data1, 1):
        for j, val in enumerate(row):
            table1.rows[i].cells[j].text = val
    
    doc.add_paragraph()
    
    # Per-Type Table
    doc.add_paragraph().add_run("Table II: Accuracy by Question Type").bold = True
    
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['Question Type', 'Supervised', 'RL']
    data2 = [
        ('Count', '82.0%', '58.0%'),
        ('Shape', '77.4%', '71.8%'),
        ('Color', '75.7%', '20.6%'),
        ('Spatial', '61.3%', '39.8%'),
    ]
    
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, row in enumerate(data2, 1):
        for j, val in enumerate(row):
            table2.rows[i].cells[j].text = val
    
    doc.add_paragraph()
    
    key_finding = doc.add_paragraph()
    key_finding.add_run("Key Finding: ").bold = True
    key_finding.add_run("RL achieves competitive performance on shape (71.8%) but struggles severely with color questions (20.6% vs. 75.7% for supervised). This suggests skill-specific composability in multimodal settings.")
    
    # ========================================
    # II. INTRODUCTION
    # ========================================
    add_section(doc, "II. INTRODUCTION")
    
    add_body_text(doc, "Compositional learning—the ability to combine existing skills into new capabilities—is fundamental to human intelligence. Recent work has demonstrated that Large Language Models can compose pretrained text skills through reinforcement learning using only sparse reward signals [1]. This raises an intriguing question: can compositional learning transfer to multimodal settings?")
    
    add_body_text(doc, "We investigate this by developing a Visual Question Answering (VQA) system that must compose: (1) Visual skills from frozen CLIP image embeddings, and (2) Language skills for question understanding and classification. Our research question is: Can RL compose frozen vision skills with trainable language skills for VQA without intermediate supervision?")
    
    add_section(doc, "A. Contributions", level=2)
    
    contributions = doc.add_paragraph()
    contributions.add_run("""Our key contributions are:
1. Systematic comparison of supervised learning vs. RL for multimodal VQA (74.0% vs. 53.7% accuracy)
2. Learning rate sensitivity analysis revealing optimal zone (1e-4 to 5e-4)
3. Per-question-type analysis showing skill-specific composability
4. Evidence that cross-modal composition is harder than text-only composition""")
    
    # ========================================
    # III. METHODOLOGY
    # ========================================
    add_section(doc, "III. METHODOLOGY")
    
    add_section(doc, "A. Model Architecture", level=2)
    
    add_body_text(doc, "Our VQA system consists of four components: (1) Vision Encoder: CLIP ViT-B/32 (frozen, 151M parameters) that converts 224×224 images into 512-dimensional embeddings; (2) Projection Layer: Trainable MLP mapping visual features to 768 dimensions; (3) Fusion Layer: Concatenation of visual and question type embeddings; (4) Classification Heads: Type-specific output heads for color (4 classes), shape (3), count (4), and spatial (13) predictions.")
    
    add_body_text(doc, "Total trainable parameters: approximately 1 million (0.6% of full model). The frozen CLIP encoder provides pretrained visual understanding while trainable components learn task-specific mappings.")
    
    add_section(doc, "B. Training Methods", level=2)
    
    add_body_text(doc, "We compare three training approaches:")
    
    add_body_text(doc, "Supervised Learning uses cross-entropy loss with ground-truth labels: L = -Σ yᵢ log(ŷᵢ). This provides direct gradient signal for every training sample.")
    
    add_body_text(doc, "REINFORCE Policy Gradient uses the update rule: ∇J(θ) = E[R · ∇θ log π(a|s;θ)] where R=1 if the predicted answer is correct, R=0 otherwise. A running mean baseline reduces variance.")
    
    add_body_text(doc, "Frozen Baseline performs no training, establishing the lower bound of random performance (approximately 4.17% for 24 classes).")
    
    add_section(doc, "C. Dataset", level=2)
    
    add_body_text(doc, "We use a synthetic CLEVR-style VQA dataset with 5,000 training samples, 1,000 validation samples, and 1,000 test samples. Each image is 224×224 pixels containing colored geometric shapes. The dataset includes four balanced question types: color (\"What color is X?\"), shape (\"What shape is X?\"), count (\"How many X?\"), and spatial (\"What is left/right of X?\"). The answer vocabulary contains 24 classes.")
    
    # ========================================
    # IV. EXPERIMENTS
    # ========================================
    add_section(doc, "IV. EXPERIMENTS")
    
    add_body_text(doc, "We conducted 61+ experiments across four categories: baseline methods, learning rate sweep, reward function variations, and question type analysis.")
    
    add_section(doc, "A. Learning Rate Sensitivity", level=2)
    
    doc.add_paragraph().add_run("Table III: RL Accuracy by Learning Rate").bold = True
    
    table3 = doc.add_table(rows=11, cols=2)
    table3.style = 'Table Grid'
    lr_data = [
        ('Learning Rate', 'Accuracy'),
        ('1e-5', '29.4%'),
        ('2e-5', '37.0%'),
        ('5e-5', '41.0%'),
        ('1e-4', '45.2%'),
        ('2e-4 (optimal)', '53.7%'),
        ('5e-4', '44.0%'),
        ('1e-3', '29.3%'),
        ('2e-3', '20.7%'),
        ('5e-3', '14.2%'),
        ('1e-2', '14.2%'),
    ]
    for i, (lr, acc) in enumerate(lr_data):
        table3.rows[i].cells[0].text = lr
        table3.rows[i].cells[1].text = acc
        if i == 0:
            table3.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            table3.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    add_body_text(doc, "The optimal learning rate is 2e-4, achieving 53.7% accuracy. Learning rates below 1e-4 result in underfitting, while rates above 1e-3 cause training collapse to near-random performance. The optimal zone spans 1e-4 to 5e-4.")
    
    add_section(doc, "B. Reward Function Comparison", level=2)
    
    add_body_text(doc, "We tested five reward functions: exact match (29.3%), partial match (29.3%), length penalty (32.4%), combined (32.4%), and progressive slow (43.1%). The progressive slow reward, which gradually increases reward difficulty, achieved the best RL performance among reward variations.")
    
    # ========================================
    # V. DISCUSSION
    # ========================================
    add_section(doc, "V. DISCUSSION")
    
    add_section(doc, "A. Why Supervised Outperforms RL", level=2)
    
    add_body_text(doc, "Three factors explain the 20+ percentage point gap between supervised (74%) and RL (53.7%) training:")
    
    add_body_text(doc, "1) Cross-modal composition challenge: Composing visual and language information across different modalities is inherently more difficult than within-modality composition as studied in text-only LLMs.")
    
    add_body_text(doc, "2) Sparse reward signal: Binary rewards (correct/incorrect) provide less gradient information per sample compared to cross-entropy loss, which penalizes every incorrect output dimension.")
    
    add_body_text(doc, "3) Training scale: With only 1000-3000 steps, RL may require 10-100× more iterations to match supervised performance, as suggested by text-only compositional learning research.")
    
    add_section(doc, "B. The Color Question Puzzle", level=2)
    
    add_body_text(doc, "RL achieves only 20.6% on color questions compared to 75.7% for supervised learning, while performing competitively on shape (71.8% vs 77.4%). This dramatic difference suggests that CLIP may encode color information more weakly than shape, or that color words in the answer vocabulary create confusion for the RL policy. The phenomenon indicates skill-specific composability in multimodal settings.")
    
    # ========================================
    # VI. LIMITATIONS
    # ========================================
    add_section(doc, "VI. LIMITATIONS")
    
    add_body_text(doc, "This study has several limitations: (1) Our synthetic CLEVR-style dataset may not reflect real-world VQA distributions; (2) CLIP ViT-B/32 (151M parameters) is relatively small compared to state-of-the-art vision models; (3) Training was limited to 1000-3000 steps due to time constraints; (4) The frozen visual encoder inherently limits spatial reasoning capability; (5) We did not explore fine-tuning the CLIP encoder or using larger vision backbones.")
    
    # ========================================
    # VII. CONCLUSION
    # ========================================
    add_section(doc, "VII. CONCLUSION")
    
    add_body_text(doc, "We investigated compositional skill learning in multimodal reinforcement learning through 61+ controlled experiments on a VQA task. Our key findings are:")
    
    conclusions = doc.add_paragraph()
    conclusions.add_run("""
1. Best accuracy: Supervised 74.0%, RL 53.7%
2. Cross-modal gap: Supervised outperforms RL by 20+ points
3. Skill-specific composability: Shape (71.8%) composes well, color (20.6%) does not
4. Optimal learning rate: 2e-4 for RL training
5. Data scaling: More data (50K) does not improve over 5K baseline

Compositional learning shows promise for multimodal systems but requires careful consideration of modality gaps, skill-specific composability, and training methodology. Future work should explore larger training budgets, fine-tuned visual encoders, and hybrid supervised-RL approaches.""")
    
    # ========================================
    # REFERENCES
    # ========================================
    add_section(doc, "REFERENCES")
    
    refs = [
        "[1] Anonymous, \"From f(x) and g(x) to f(g(x)): LLMs Learn New Skills in RL by Composing Old Ones,\" arXiv:2509.25123, 2024.",
        "[2] A. Radford et al., \"Learning Transferable Visual Models From Natural Language Supervision,\" ICML, 2021.",
        "[3] R. J. Williams, \"Simple Statistical Gradient-Following Algorithms for Connectionist Reinforcement Learning,\" Machine Learning, 1992.",
        "[4] S. Antol et al., \"VQA: Visual Question Answering,\" ICCV, 2015.",
        "[5] J. Li et al., \"BLIP-2: Bootstrapping Language-Image Pre-training,\" ICML, 2023.",
        "[6] J. Johnson et al., \"CLEVR: A Diagnostic Dataset for Compositional Language and Elementary Visual Reasoning,\" CVPR, 2017.",
        "[7] B. M. Lake et al., \"Building Machines That Learn and Think Like People,\" Behavioral and Brain Sciences, 2017.",
        "[8] C. Keysers et al., \"Measuring Compositional Generalization,\" ICLR, 2020.",
    ]
    
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(9)
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, "IEEE_PROJECT_REPORT.docx")
    doc.save(output_path)
    print(f"IEEE Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    print("=" * 60)
    print("GENERATING IEEE FORMAT PROJECT REPORT")
    print("=" * 60)
    
    output = create_ieee_report()
    
    print("\n" + "=" * 60)
    print("REPORT COMPLETE")
    print(f"Output: {output}")
    print("=" * 60)
